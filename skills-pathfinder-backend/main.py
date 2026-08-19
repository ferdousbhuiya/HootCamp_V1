# ============================================================
# Skills Pathfinder Backend
# Version 1.2.0
#
# Step 4B:
# Comprehensive document extraction and skill analysis
# ============================================================

import os
import io
import json
import re
import traceback
from typing import List, Dict, Any, Optional

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware

from dotenv import load_dotenv
from supabase import create_client, Client

from PyPDF2 import PdfReader
import pytesseract
from PIL import Image, ImageOps, ImageEnhance, ImageFilter
import docx

from groq import Groq

from pydantic import BaseModel

from recommendation_engine import (
    get_career_recommendations,
    get_skill_gap_analysis,
)

# Optional scanned-PDF OCR support
try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


# ============================================================
# 1. ENVIRONMENT
# ============================================================

load_dotenv()


# ============================================================
# 2. FASTAPI APPLICATION
# ============================================================

app = FastAPI(
    title="Skills Pathfinder API",
    version="1.2.0",
    description=(
        "Backend API for Skills Pathfinder. "
        "Supports document extraction, OCR, comprehensive "
        "skill analysis, career recommendations and career reports."
    ),
)


# ============================================================
# 3. CORS
# ============================================================

ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:3000,"
    "http://localhost:5173,"
    "https://ferdouse.us,"
    "https://www.ferdouse.us,"
    "https://skillpathfinder.ferdous.us",
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in ALLOWED_ORIGINS],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# 4. SUPABASE
# ============================================================

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

_supabase_client: Optional[Client] = None


def get_supabase() -> Client:
    """
    Lazily initialize Supabase.

    The current document-processing endpoints do not require
    Supabase. This prevents a Supabase configuration problem
    from stopping the API from starting.
    """
    global _supabase_client

    if _supabase_client is None:
        if not SUPABASE_URL or not SUPABASE_KEY:
            raise RuntimeError(
                "SUPABASE_URL and SUPABASE_KEY must be configured."
            )

        _supabase_client = create_client(
            SUPABASE_URL,
            SUPABASE_KEY,
        )

    return _supabase_client


# ============================================================
# 5. GROQ / LLM
# ============================================================

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

GROQ_MODEL = os.getenv(
    "GROQ_MODEL",
    "openai/gpt-oss-20b",
)

if not GROQ_API_KEY:
    print("[WARN] GROQ_API_KEY is not configured.")

groq_client = Groq(api_key=GROQ_API_KEY)


async def llm_generate(
    prompt: str,
    max_tokens_override: int = 8192,
) -> str:
    """
    Send a prompt to Groq and return the JSON response.
    """

    try:
        print(
            f"[LLM] Calling Groq API "
            f"model={GROQ_MODEL}, "
            f"max_tokens={max_tokens_override}"
        )

        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert resume parser, technical "
                        "recruiter, career analyst and skills taxonomy "
                        "specialist. "
                        "Return ONLY valid JSON. "
                        "Never return markdown. "
                        "Never add explanations outside the JSON object."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.05,
            max_tokens=max_tokens_override,
            response_format={"type": "json_object"},
        )

        response_text = completion.choices[0].message.content or ""

        print(
            f"[LLM] Response length: "
            f"{len(response_text)} characters"
        )

        return response_text

    except Exception as e:
        print(f"[LLM] Groq API Error: {str(e)}")

        raise HTTPException(
            status_code=502,
            detail=f"Failed to connect to Groq API: {str(e)}",
        )


# ============================================================
# 6. TEXT CLEANING
# ============================================================

def clean_extracted_text(text: str) -> str:
    """
    Clean OCR/PDF/DOCX text without destroying useful information.
    """

    if not text:
        return ""

    # Normalize different line endings
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Remove null characters
    text = text.replace("\x00", "")

    # Normalize tabs
    text = text.replace("\t", " ")

    # Remove excessive spaces
    text = re.sub(r"[ ]{2,}", " ", text)

    # Preserve lines but remove excessive blank lines
    text = re.sub(r"\n[ ]+\n", "\n\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ============================================================
# 7. IMAGE PREPROCESSING FOR OCR
# ============================================================

def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """
    Improve OCR quality for photographs and scanned documents.
    """

    # Correct orientation where possible
    try:
        image = ImageOps.exif_transpose(image)
    except Exception:
        pass

    # Convert to grayscale
    image = image.convert("L")

    # Increase contrast
    image = ImageEnhance.Contrast(image).enhance(1.7)

    # Slight sharpening
    image = image.filter(ImageFilter.SHARPEN)

    return image


# ============================================================
# 8. PDF EXTRACTION
# ============================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract selectable text from PDF.

    If very little text is found, automatically use OCR.
    """

    extracted_parts = []

    # --------------------------------------------------------
    # First attempt: normal PDF text extraction
    # --------------------------------------------------------

    try:
        reader = PdfReader(io.BytesIO(file_bytes))

        print(
            f"[PDF] Number of pages: "
            f"{len(reader.pages)}"
        )

        for page_number, page in enumerate(reader.pages, start=1):

            try:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    extracted_parts.append(
                        f"[Page {page_number}]\n{page_text}"
                    )

            except Exception as page_error:
                print(
                    f"[PDF] Error extracting page "
                    f"{page_number}: {page_error}"
                )

    except Exception as e:
        print(f"[PDF] Standard extraction error: {e}")

    text = clean_extracted_text(
        "\n\n".join(extracted_parts)
    )

    print(
        f"[PDF] Standard extraction: "
        f"{len(text)} characters"
    )

    # --------------------------------------------------------
    # OCR fallback
    # --------------------------------------------------------

    # If less than 100 meaningful characters were extracted,
    # assume the PDF may be scanned.
    if len(text.strip()) < 100:

        if not PDF2IMAGE_AVAILABLE:
            print(
                "[PDF] pdf2image unavailable. "
                "Cannot OCR scanned PDF."
            )

            return text

        print(
            "[PDF] Very little selectable text found. "
            "Starting OCR..."
        )

        try:

            images = convert_from_bytes(
                file_bytes,
                dpi=250,
                fmt="png",
            )

            ocr_parts = []

            for page_number, image in enumerate(
                images,
                start=1,
            ):

                print(
                    f"[OCR] Processing PDF page "
                    f"{page_number}/{len(images)}"
                )

                processed_image = preprocess_image_for_ocr(
                    image
                )

                page_text = pytesseract.image_to_string(
                    processed_image,
                    config="--psm 6",
                )

                if page_text.strip():
                    ocr_parts.append(
                        f"[Page {page_number}]\n{page_text}"
                    )

            ocr_text = clean_extracted_text(
                "\n\n".join(ocr_parts)
            )

            print(
                f"[OCR] PDF OCR extracted "
                f"{len(ocr_text)} characters"
            )

            # Prefer OCR if it found more information
            if len(ocr_text) > len(text):
                text = ocr_text

        except Exception as e:
            print(
                f"[OCR] Scanned PDF OCR failed: {e}"
            )

    return clean_extracted_text(text)


# ============================================================
# 9. DOCX EXTRACTION
# ============================================================

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX paragraphs AND tables.
    """

    try:

        document = docx.Document(
            io.BytesIO(file_bytes)
        )

        parts = []

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):

            parts.append(
                f"[Table {table_index}]"
            )

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        cells.append(cell_text)

                if cells:
                    parts.append(
                        " | ".join(cells)
                    )

        text = clean_extracted_text(
            "\n".join(parts)
        )

        print(
            f"[DOCX] Extracted "
            f"{len(text)} characters"
        )

        return text

    except Exception as e:

        print(
            f"[DOCX] Extraction error: {e}"
        )

        raise HTTPException(
            status_code=400,
            detail="Invalid or corrupted DOCX file.",
        )


# ============================================================
# 10. IMAGE OCR
# ============================================================

def extract_text_from_image(
    file_bytes: bytes,
) -> str:
    """
    OCR JPG, JPEG and PNG images.
    """

    try:

        image = Image.open(
            io.BytesIO(file_bytes)
        )

        print(
            f"[IMAGE] Original image size: "
            f"{image.size}"
        )

        processed_image = preprocess_image_for_ocr(
            image
        )

        # PSM 6 works well for blocks of text such as resumes.
        text = pytesseract.image_to_string(
            processed_image,
            config="--psm 6",
        )

        text = clean_extracted_text(text)

        print(
            f"[OCR] Image extracted "
            f"{len(text)} characters"
        )

        return text

    except Exception as e:

        print(
            f"[IMAGE] OCR error: {e}"
        )

        raise HTTPException(
            status_code=400,
            detail="Invalid or unreadable image file.",
        )


# ============================================================
# 11. TXT EXTRACTION
# ============================================================

def extract_text_from_txt(
    file_bytes: bytes,
) -> str:
    """
    Extract UTF-8 text, with fallback encodings.
    """

    encodings = [
        "utf-8",
        "utf-8-sig",
        "latin-1",
    ]

    for encoding in encodings:

        try:

            text = file_bytes.decode(
                encoding
            )

            return clean_extracted_text(text)

        except UnicodeDecodeError:
            continue

    raise HTTPException(
        status_code=400,
        detail="Unable to decode text file.",
    )


# ============================================================
# 12. DOCUMENT EXTRACTION ROUTER
# ============================================================

def extract_document_text(
    file_bytes: bytes,
    file_ext: str,
) -> str:

    if file_ext == ".pdf":
        return extract_text_from_pdf(
            file_bytes
        )

    if file_ext == ".docx":
        return extract_text_from_docx(
            file_bytes
        )

    if file_ext == ".txt":
        return extract_text_from_txt(
            file_bytes
        )

    if file_ext in {
        ".png",
        ".jpg",
        ".jpeg",
    }:
        return extract_text_from_image(
            file_bytes
        )

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type.",
    )


# ============================================================
# 13. SKILL NORMALIZATION
# ============================================================

def normalize_skill_name(
    skill_name: str,
) -> str:
    """
    Normalize common formatting variations without
    changing the actual skill meaning.
    """

    if not skill_name:
        return ""

    name = skill_name.strip()

    replacements = {
        "ms excel": "Microsoft Excel",
        "microsoft excel": "Microsoft Excel",
        "ms word": "Microsoft Word",
        "microsoft word": "Microsoft Word",
        "powerbi": "Power BI",
        "power bi": "Power BI",
        "sql server": "Microsoft SQL Server",
        "postgres": "PostgreSQL",
        "postgresql": "PostgreSQL",
        "js": "JavaScript",
        "nodejs": "Node.js",
        "node js": "Node.js",
        "reactjs": "React",
        "react.js": "React",
        "py": "Python",
        "machine learning": "Machine Learning",
        "artificial intelligence": "Artificial Intelligence",
    }

    key = name.lower()

    if key in replacements:
        return replacements[key]

    return name


def deduplicate_skills(
    skills: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    Remove duplicate skills while retaining the strongest
    available evidence/confidence.
    """

    merged = {}

    for skill in skills:

        if not isinstance(skill, dict):
            continue

        name = normalize_skill_name(
            str(skill.get("name", ""))
        )

        if not name:
            continue

        key = name.lower()

        skill["name"] = name

        existing = merged.get(key)

        if existing is None:
            merged[key] = skill
            continue

        # Prefer higher confidence
        current_confidence = float(
            existing.get("confidence", 0)
            or 0
        )

        new_confidence = float(
            skill.get("confidence", 0)
            or 0
        )

        if new_confidence > current_confidence:
            merged[key] = skill

        else:

            # Preserve evidence from both occurrences
            old_evidence = existing.get(
                "evidence",
                "",
            )

            new_evidence = skill.get(
                "evidence",
                "",
            )

            if new_evidence and new_evidence not in old_evidence:
                existing["evidence"] = (
                    f"{old_evidence}; {new_evidence}"
                    if old_evidence
                    else new_evidence
                )

    return list(merged.values())


# ============================================================
# 14. COMPREHENSIVE AI EXTRACTION
# ============================================================

async def extract_skills_with_llm(
    text: str,
) -> dict:

    # Protect the API from pathological input sizes.
    # We keep a large amount of resume content while preventing
    # extremely large uploads from consuming the entire context.
    MAX_TEXT_CHARS = 60000

    full_text = text[:MAX_TEXT_CHARS]

    prompt = f"""
You are an expert technical recruiter, resume parser,
career counselor and skills taxonomy specialist.

Analyze the COMPLETE document below.

Your goal is to build a comprehensive professional profile
from the document.

DO NOT return only the skills from a "Skills" section.

You MUST inspect ALL available sections including:

- Professional Summary
- Objective
- Skills
- Technical Skills
- Work Experience
- Employment
- Education
- Degrees
- Courses
- Certifications
- Training
- Projects
- Internships
- Research
- Publications
- Volunteer Experience
- Achievements
- Responsibilities
- Job descriptions
- Technologies mentioned inside sentences

============================================================
SKILL EXTRACTION
============================================================

Extract EVERY meaningful skill that is explicitly supported
by the document.

Do NOT unnecessarily group skills.

For example, if the document says:

"Used Python, Pandas, NumPy, SQL and Power BI"

extract:

Python
Pandas
NumPy
SQL
Power BI

Do NOT return only:

Data Analytics

However, broader domain skills should ALSO be extracted when
explicitly supported.

Examples:

Python
Pandas
NumPy
SQL
Data Analysis

============================================================
SKILL CATEGORIES
============================================================

Use one of these categories whenever possible:

Programming Language
Framework/Library
Tool/Software
Database
Cloud Technology
Data/Analytics
AI/Machine Learning
Engineering
Domain Knowledge
Methodology/Standard
Project Management
Business
Communication
Leadership
Soft Skill
Certification
Other Technical Skill

============================================================
DO NOT INVENT
============================================================

Only extract information supported by the document.

Do not assume a person knows a technology simply because
another technology normally requires it.

Do not infer Python just because someone is a data analyst.

Do not infer AWS just because someone worked in cloud computing.

============================================================
PROFICIENCY
============================================================

Determine proficiency only when evidence exists.

Allowed values:

beginner
intermediate
advanced
expert
unknown

If the document does not provide enough evidence,
use "unknown".

============================================================
CONFIDENCE
============================================================

Use a number between 0 and 1.

0.95-1.00:
Very clearly stated.

0.80-0.94:
Strong evidence.

0.60-0.79:
Reasonable explicit evidence.

Below 0.60:
Use only if there is still explicit supporting evidence.

Do not invent skills merely to increase the number.

============================================================
EVIDENCE
============================================================

For every skill, provide a short evidence string showing
WHY the skill was extracted.

Example:

"Used Python and Pandas for data analysis projects."

============================================================
CERTIFICATIONS
============================================================

Extract every certification mentioned.

For each certification identify:

name
provider
credential_id
verification_url
status

============================================================
EDUCATION
============================================================

Extract:

degree
institution
field_of_study
graduation_year

============================================================
WORK EXPERIENCE
============================================================

Extract:

job_title
company
start_date
end_date
responsibilities
technologies_used

============================================================
PROJECTS
============================================================

Extract:

project_name
description
technologies
skills_demonstrated

============================================================
ONGOING LEARNING
============================================================

Identify courses, training programs and learning activities
that appear to be ongoing.

============================================================
OUTPUT
============================================================

Return ONLY valid JSON.

Use EXACTLY this structure:

{{
  "document_type": "resume|cv|certificate|course_document|other",

  "candidate_profile": {{
    "name": "",
    "professional_summary": ""
  }},

  "education": [
    {{
      "degree": "",
      "institution": "",
      "field_of_study": "",
      "graduation_year": ""
    }}
  ],

  "work_experience": [
    {{
      "job_title": "",
      "company": "",
      "start_date": "",
      "end_date": "",
      "responsibilities": [],
      "technologies_used": []
    }}
  ],

  "projects": [
    {{
      "project_name": "",
      "description": "",
      "technologies": [],
      "skills_demonstrated": []
    }}
  ],

  "certifications": [
    {{
      "name": "",
      "provider": "",
      "credential_id": "",
      "verification_url": "",
      "status": ""
    }}
  ],

  "courses_and_training": [
    {{
      "name": "",
      "provider": "",
      "status": "",
      "completion_date": ""
    }}
  ],

  "skills": [
    {{
      "name": "",
      "category": "",
      "proficiency_level": "beginner|intermediate|advanced|expert|unknown",
      "confidence": 0.0,
      "evidence": ""
    }}
  ],

  "key_strengths": [],

  "skill_gaps_explicitly_mentioned": [],

  "keywords": []
}}

IMPORTANT:

1. Extract all explicitly supported skills.
2. Preserve specific technologies.
3. Do not collapse specific technologies into broad categories.
4. Do not invent unsupported skills.
5. Return valid JSON only.

DOCUMENT:

{full_text}
"""

    try:

        llm_response = await llm_generate(
            prompt,
            max_tokens_override=8192,
        )

        data = json.loads(
            llm_response
        )

        if not isinstance(data, dict):
            raise ValueError(
                "LLM returned JSON that is not an object."
            )

        skills = data.get(
            "skills",
            [],
        )

        if not isinstance(skills, list):
            skills = []

        skills = deduplicate_skills(
            skills
        )

        data["skills"] = skills

        print(
            f"[SKILLS] Successfully extracted "
            f"{len(skills)} unique skills"
        )

        return data

    except json.JSONDecodeError as e:

        print(
            f"[ERROR] JSON Decode Error: {e}"
        )

        print(
            f"[ERROR] Raw LLM response: "
            f"{llm_response[:2000] if 'llm_response' in locals() else 'N/A'}"
        )

        return {
            "document_type": "other",
            "candidate_profile": {},
            "education": [],
            "work_experience": [],
            "projects": [],
            "certifications": [],
            "courses_and_training": [],
            "skills": [],
            "key_strengths": [],
            "skill_gaps_explicitly_mentioned": [],
            "keywords": [],
        }

    except Exception as e:

        print(
            f"[ERROR] LLM skill extraction failed: "
            f"{str(e)}"
        )

        traceback.print_exc()

        return {
            "document_type": "other",
            "candidate_profile": {},
            "education": [],
            "work_experience": [],
            "projects": [],
            "certifications": [],
            "courses_and_training": [],
            "skills": [],
            "key_strengths": [],
            "skill_gaps_explicitly_mentioned": [],
            "keywords": [],
        }


# ============================================================
# 15. ROOT
# ============================================================

@app.get("/")
def read_root():

    return {
        "status": "Skills Pathfinder API is online.",
        "version": "1.2.0",
    }


# ============================================================
# 16. HEALTH / ENVIRONMENT
# ============================================================

@app.get("/debug-env")
def debug_env():

    return {
        "groq_key_set": bool(
            os.getenv("GROQ_API_KEY")
        ),

        "groq_model": GROQ_MODEL,

        "supabase_configured": bool(
            os.getenv("SUPABASE_URL")
            and os.getenv("SUPABASE_KEY")
        ),

        "tesseract_command": (
            pytesseract.pytesseract.tesseract_cmd
            if pytesseract.pytesseract.tesseract_cmd
            else "default"
        ),

        "pdf2image_available": PDF2IMAGE_AVAILABLE,
    }


# ============================================================
# 17. DOCUMENT UPLOAD
# ============================================================

@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(...),
):

    filename = (
        file.filename
        or "unknown"
    )

    print(
        f"[UPLOAD] Received file: "
        f"{filename}"
    )

    try:

        # ----------------------------------------------------
        # Validate extension
        # ----------------------------------------------------

        allowed_extensions = {
            ".pdf",
            ".docx",
            ".png",
            ".jpeg",
            ".jpg",
            ".txt",
        }

        file_ext = os.path.splitext(
            filename
        )[1].lower()

        if file_ext not in allowed_extensions:

            raise HTTPException(
                status_code=400,
                detail=(
                    "Unsupported file type. "
                    "Allowed: PDF, DOCX, TXT, "
                    "PNG, JPG, JPEG."
                ),
            )

        # ----------------------------------------------------
        # Read file
        # ----------------------------------------------------

        file_bytes = await file.read()

        file_size = len(file_bytes)

        print(
            f"[UPLOAD] File size: "
            f"{file_size:,} bytes"
        )

        # 15 MB limit
        MAX_FILE_SIZE = 15 * 1024 * 1024

        if file_size > MAX_FILE_SIZE:

            raise HTTPException(
                status_code=400,
                detail=(
                    "File size exceeds "
                    "15MB limit."
                ),
            )

        if file_size == 0:

            raise HTTPException(
                status_code=400,
                detail="Uploaded file is empty.",
            )

        # ----------------------------------------------------
        # Extract text
        # ----------------------------------------------------

        print(
            f"[EXTRACT] Processing "
            f"{file_ext} file..."
        )

        text = extract_document_text(
            file_bytes,
            file_ext,
        )

        text = clean_extracted_text(
            text
        )

        character_count = len(text)

        print(
            f"[EXTRACT] Final text length: "
            f"{character_count} characters"
        )

        if character_count < 10:

            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract meaningful text "
                    "from this document. "
                    "The file may be blank, corrupted, "
                    "or contain text that OCR cannot read."
                ),
            )

        # ----------------------------------------------------
        # AI extraction
        # ----------------------------------------------------

        print(
            "[AI] Starting comprehensive "
            "document analysis..."
        )

        analysis = await extract_skills_with_llm(
            text
        )

        skills = analysis.get(
            "skills",
            [],
        )

        # ----------------------------------------------------
        # Career recommendations
        # ----------------------------------------------------

        try:

            recommendations = (
                get_career_recommendations(
                    skills,
                    top_n=5,
                )
            )

        except Exception as recommendation_error:

            print(
                "[WARN] Career recommendation "
                f"error: {recommendation_error}"
            )

            recommendations = []

        # ----------------------------------------------------
        # Final response
        # ----------------------------------------------------

        response = {

            "message": (
                "File processed successfully"
            ),

            "filename": filename,

            "file_type": file_ext,

            "character_count": character_count,

            "extracted_skills": skills,

            "skills_count": len(skills),

            "explanations": analysis.get(
                "skills",
                [],
            ),

            "recommendations": recommendations,

            "document_type": analysis.get(
                "document_type",
                "other",
            ),

            "candidate_profile": analysis.get(
                "candidate_profile",
                {},
            ),

            "education": analysis.get(
                "education",
                [],
            ),

            "work_experience": analysis.get(
                "work_experience",
                [],
            ),

            "projects": analysis.get(
                "projects",
                [],
            ),

            "certifications": analysis.get(
                "certifications",
                [],
            ),

            "courses_and_training": analysis.get(
                "courses_and_training",
                [],
            ),

            "key_strengths": analysis.get(
                "key_strengths",
                [],
            ),

            "skill_gaps_explicitly_mentioned": analysis.get(
                "skill_gaps_explicitly_mentioned",
                [],
            ),

            "keywords": analysis.get(
                "keywords",
                [],
            ),
        }

        print(
            f"[SUCCESS] {filename}: "
            f"{character_count} characters, "
            f"{len(skills)} skills"
        )

        return response

    except HTTPException:
        raise

    except Exception as e:

        traceback.print_exc()

        print(
            f"[ERROR] Unexpected upload error: "
            f"{str(e)}"
        )

        return JSONResponse(
            status_code=500,
            content={
                "detail": (
                    "Server error processing file: "
                    f"{str(e)}"
                ),
                "partial": True,
            },
        )


# ============================================================
# 18. CERTIFICATE VERIFICATION
# ============================================================

@app.post("/api/verify-certificate")
async def verify_certificate(
    file: UploadFile = File(...),
):

    filename = (
        file.filename
        or "certificate"
    )

    try:

        file_bytes = await file.read()

        file_ext = os.path.splitext(
            filename
        )[1].lower()

        if file_ext not in {
            ".pdf",
            ".png",
            ".jpg",
            ".jpeg",
        }:

            raise HTTPException(
                status_code=400,
                detail=(
                    "Unsupported certificate file. "
                    "Use PDF, PNG, JPG or JPEG."
                ),
            )

        text = extract_document_text(
            file_bytes,
            file_ext,
        )

        if not text.strip():

            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract text "
                    "from certificate."
                ),
            )

        prompt = f"""
Extract certificate information from the following text.

Return ONLY valid JSON.

Required structure:

{{
  "certification_name": "",
  "provider": "",
  "credential_id": "",
  "verification_url": "",
  "certificate_status": ""
}}

Rules:

- Do not invent information.
- If a field is not present, return an empty string.
- Preserve credential IDs exactly.
- Preserve URLs exactly when possible.

Certificate Text:

{text}
"""

        llm_response = await llm_generate(
            prompt,
            max_tokens_override=2048,
        )

        cert_data = json.loads(
            llm_response
        )

        if not cert_data.get(
            "certification_name"
        ):

            cert_data[
                "certification_name"
            ] = os.path.splitext(
                filename
            )[0]

        if not cert_data.get(
            "provider"
        ):

            cert_data[
                "provider"
            ] = "Unknown"

        # ----------------------------------------------------
        # Automatic URL recognition
        # ----------------------------------------------------

        verification_url = (
            cert_data.get(
                "verification_url",
                "",
            )
            or ""
        )

        if verification_url:

            known_domains = (
                r"(aws\.amazon\.com|"
                r"coursera\.org|"
                r"udemy\.com|"
                r"edx\.org|"
                r"google\.com|"
                r"microsoft\.com|"
                r"cisco\.com|"
                r"comptia\.org|"
                r"credly\.com|"
                r"credential\.net)"
            )

            cert_data[
                "auto_verified"
            ] = bool(
                re.search(
                    known_domains,
                    verification_url,
                    re.IGNORECASE,
                )
            )

        else:

            cert_data[
                "auto_verified"
            ] = False

        return cert_data

    except HTTPException:
        raise

    except Exception as e:

        traceback.print_exc()

        raise HTTPException(
            status_code=500,
            detail=(
                "Error processing certificate: "
                f"{str(e)}"
            ),
        )


# ============================================================
# 19. SKILL REQUEST MODEL
# ============================================================

class SkillsRequest(BaseModel):

    extracted_skills: List[
        Dict[str, Any]
    ]


# ============================================================
# 20. CAREER RECOMMENDATIONS
# ============================================================

@app.post("/api/recommendations")
async def get_recommendations(
    skills_request: SkillsRequest,
):

    try:

        recommendations = (
            get_career_recommendations(
                skills_request.extracted_skills,
                top_n=5,
            )
        )

        return {
            "status": "success",
            "total_recommendations": len(
                recommendations
            ),
            "recommendations": recommendations,
        }

    except HTTPException:
        raise

    except Exception as e:

        traceback.print_exc()

        raise HTTPException(
            status_code=500,
            detail=(
                "Error generating recommendations: "
                f"{str(e)}"
            ),
        )


# ============================================================
# 21. CAREER PATHS
# ============================================================

@app.get("/api/career-paths")
async def get_all_career_paths():

    from recommendation_engine import (
        CAREER_PATHS,
    )

    career_list = [

        {
            "id": career["id"],
            "path": career["path"],
            "category": career["category"],
            "job_outlook": career[
                "job_outlook"
            ],
            "median_salary": career[
                "median_salary"
            ],
        }

        for career in CAREER_PATHS
    ]

    return {
        "status": "success",
        "career_paths": career_list,
    }


# ============================================================
# 22. SKILL GAP ANALYSIS
# ============================================================

@app.post("/api/skill-gap-analysis")
async def skill_gap_analysis(
    skills_request: SkillsRequest,
    career_id: str,
):

    try:

        analysis = (
            get_skill_gap_analysis(
                skills_request.extracted_skills,
                career_id,
            )
        )

        if not analysis:

            raise HTTPException(
                status_code=404,
                detail="Career path not found",
            )

        return {
            "status": "success",
            "analysis": analysis,
        }

    except HTTPException:
        raise

    except Exception as e:

        print(
            f"[ERROR] Skill gap analysis: "
            f"{str(e)}"
        )

        raise HTTPException(
            status_code=500,
            detail=(
                "Error analyzing skill gap: "
                f"{str(e)}"
            ),
        )


# ============================================================
# 23. CAREER REPORT
# ============================================================

class ReportRequest(BaseModel):

    skills: List[str]


@app.post("/api/generate-career-advice")
async def generate_career_advice(
    request: ReportRequest,
):

    try:

        skills_text = (
            ", ".join(request.skills)
            if request.skills
            else "No skills provided"
        )

        prompt = f"""
You are an expert Career Coach and Talent Analyst.

Based on the following skills:

[{skills_text}]

Generate a comprehensive career report.

Return ONLY valid JSON.

Required structure:

{{
  "executive_summary": "",
  "swot_analysis": {{
    "strengths": [],
    "weaknesses": [],
    "opportunities": [],
    "threats": []
  }},
  "action_plan": {{
    "30_days": "",
    "60_days": "",
    "90_days": "",
    "6_months": "",
    "1_year": ""
  }},
  "recommended_next_skills": []
}}
"""

        llm_response = await llm_generate(
            prompt,
            max_tokens_override=4096,
        )

        advice_data = json.loads(
            llm_response
        )

        return {
            "status": "success",
            "advice": advice_data,
        }

    except HTTPException:
        raise

    except Exception as e:

        traceback.print_exc()

        raise HTTPException(
            status_code=500,
            detail=(
                "Error generating career advice: "
                f"{str(e)}"
            ),
        )


# ============================================================
# 24. STARTUP INFORMATION
# ============================================================

@app.on_event("startup")
async def startup_event():

    print("=" * 60)
    print("Skills Pathfinder Backend")
    print("Version: 1.2.0")
    print("=" * 60)

    print(
        f"Tesseract: "
        f"{pytesseract.pytesseract.tesseract_cmd}"
    )

    print(
        f"PDF OCR available: "
        f"{PDF2IMAGE_AVAILABLE}"
    )

    print(
        f"Groq model: "
        f"{GROQ_MODEL}"
    )

    print(
        f"Supabase configured: "
        f"{bool(SUPABASE_URL and SUPABASE_KEY)}"
    )

    print("=" * 60)
